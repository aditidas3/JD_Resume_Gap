"""
resume_reader.py

Reads resume text out of a real file so you don't have to paste it as a
Python string. Supports .docx and .pdf -- picks the right library based on
the file extension.
"""

from pathlib import Path


def read_resume(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _read_docx(path)
    elif suffix == ".pdf":
        return _read_pdf(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8")
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Use .docx, .pdf, or .txt, "
            f"or convert your resume to one of these first."
        )


def _read_docx(path: Path) -> str:
    from docx import Document  # python-docx package (imports as `docx`)

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Resumes often have skills/details inside tables, not just paragraphs --
    # don't silently drop that content.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages_text)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python resume_reader.py <path-to-resume>")
    else:
        text = read_resume(sys.argv[1])
        print(f"--- Extracted {len(text)} characters ---\n")
        print(text[:1000])
        print("\n... (truncated)" if len(text) > 1000 else "")
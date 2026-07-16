"""
docx_writer.py

Edits an EXISTING resume.docx IN PLACE: replaces the professional summary
text and inserts new bullets directly into the relevant existing sections
(matched by content similarity), rather than appending a separate
"suggestions" section at the end. All original formatting (fonts, bullet
style, spacing) is preserved because new bullets are created by cloning an
existing bullet paragraph's XML and swapping only its text.

"unaddressable_gaps" are NOT written into the resume (a resume shouldn't
list its own shortcomings) -- they're returned separately so you can log or
review them outside the document.
"""

import copy
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor

HIGHLIGHT_COLOR = RGBColor(0x1F, 0x4E, 0x9C)  # blue

STOPWORDS = {
    "the", "a", "an", "and", "or", "of", "to", "in", "on", "for", "with",
    "is", "are", "was", "were", "this", "that", "as", "by", "at", "from",
}


def _words(text: str) -> set:
    return {w for w in re.findall(r"[a-zA-Z0-9]+", text.lower()) if w not in STOPWORDS}


def _replace_paragraph_text(paragraph: Paragraph, new_text: str, highlight: bool = False) -> None:
    """Overwrite a paragraph's visible text while keeping its first run's
    formatting (font, size, bold, etc.) -- clears any extra runs so no
    fragments of the old text linger. If highlight=True, marks the new text
    italic + blue so it's visually obvious what the tool changed -- remove
    this formatting (select text -> clear formatting) before sending the
    resume out for real."""
    if not paragraph.runs:
        run = paragraph.add_run(new_text)
        if highlight:
            run.italic = True
            run.font.color.rgb = HIGHLIGHT_COLOR
        return
    paragraph.runs[0].text = new_text
    if highlight:
        paragraph.runs[0].italic = True
        paragraph.runs[0].font.color.rgb = HIGHLIGHT_COLOR
    for run in paragraph.runs[1:]:
        run.text = ""


def _insert_paragraph_after(anchor: Paragraph, new_text: str, highlight: bool = True) -> Paragraph:
    """Clone `anchor`'s paragraph (same style/bullet formatting), insert it
    directly after `anchor` in the document, and set its text to `new_text`.
    This is how new bullets end up looking identical to your existing ones.
    highlight=True by default here since every bullet this function inserts
    is, by definition, new content."""
    new_p_element = copy.deepcopy(anchor._p)
    anchor._p.addnext(new_p_element)
    new_paragraph = Paragraph(new_p_element, anchor._parent)
    _replace_paragraph_text(new_paragraph, new_text, highlight=highlight)
    return new_paragraph


def _find_summary_paragraph(doc: Document) -> Paragraph | None:
    """Find the paragraph right after a 'PROFESSIONAL SUMMARY' heading."""
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        if "professional summary" in p.text.strip().lower():
            # the summary content is typically the next non-empty paragraph
            for j in range(i + 1, len(paragraphs)):
                if paragraphs[j].text.strip():
                    return paragraphs[j]
    return None


def _find_best_matching_bullet(doc: Document, query_text: str) -> Paragraph | None:
    """Find the existing bullet paragraph whose text has the highest word
    overlap with query_text (typically the 'based_on' field from a
    suggested bullet) -- this is the anchor we insert the new bullet after."""
    query_words = _words(query_text)
    if not query_words:
        return None

    best_paragraph = None
    best_score = 0
    for p in doc.paragraphs:
        if p.style.name != "List Paragraph" or not p.text.strip():
            continue
        overlap = len(query_words & _words(p.text))
        if overlap > best_score:
            best_score = overlap
            best_paragraph = p

    return best_paragraph if best_score > 0 else None


def write_tailored_docx(original_resume_path: str, draft: dict, output_path: str) -> dict:
    """
    Returns a dict with:
      - output_path: where the tailored .docx was saved
      - bullets_placed: bullets successfully inserted into a matched section
      - bullets_skipped: bullets with no good match found (not inserted --
        logged here so you can add them manually if you want them)
      - unaddressable_gaps: passed through from draft, for your own review
        (never written into the resume itself)
    """
    doc = Document(original_resume_path)

    # 1. Replace the professional summary in place -- highlighted since it's
    # entirely new/changed content (per your call: summary gets highlighted too)
    summary_paragraph = _find_summary_paragraph(doc)
    if summary_paragraph is not None:
        _replace_paragraph_text(summary_paragraph, draft.get("tailored_summary", ""), highlight=True)
        summary_updated = True
    else:
        summary_updated = False

    # 2. Insert each suggested bullet directly after its best-matching
    # existing bullet, so it lands in the right section with matching
    # formatting -- instead of a separate appended section.
    bullets_placed = []
    bullets_skipped = []
    for b in draft.get("suggested_bullets", []):
        anchor = _find_best_matching_bullet(doc, b.get("based_on", ""))
        if anchor is not None:
            _insert_paragraph_after(anchor, b["bullet_text"])
            bullets_placed.append(b["bullet_text"])
        else:
            bullets_skipped.append(b["bullet_text"])

    doc.save(output_path)

    return {
        "output_path": output_path,
        "summary_updated": summary_updated,
        "bullets_placed": bullets_placed,
        "bullets_skipped": bullets_skipped,
        "unaddressable_gaps": draft.get("unaddressable_gaps", []),
    }


if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 4:
        print("Usage: python docx_writer.py <original_resume.docx> <draft_json_file> <output.docx>")
        sys.exit(1)

    with open(sys.argv[2]) as f:
        draft = json.load(f)

    result = write_tailored_docx(sys.argv[1], draft, sys.argv[3])
    print(f"Saved to: {result['output_path']}")
    print(f"Summary updated: {result['summary_updated']}")
    print(f"Bullets placed: {len(result['bullets_placed'])}")
    for bt in result["bullets_placed"]:
        print(f"  + {bt}")
    if result["bullets_skipped"]:
        print(f"Bullets skipped (no good match found): {len(result['bullets_skipped'])}")
        for bt in result["bullets_skipped"]:
            print(f"  ? {bt}")
    if result["unaddressable_gaps"]:
        print("Unaddressable gaps (not written to resume -- for your own review):")
        for g in result["unaddressable_gaps"]:
            print(f"  - {g}")